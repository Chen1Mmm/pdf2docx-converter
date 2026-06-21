"""
PDF转Word转换器 - 支持批量处理和扫描版PDF
支持OCR识别扫描版PDF
"""

import os
import sys
import tempfile
import urllib.request
from pathlib import Path
from typing import List, Tuple
import threading
from datetime import datetime

import pytesseract
from pdf2image import convert_from_path
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QFileDialog, QListWidget, QListWidgetItem, QLabel,
    QProgressBar, QCheckBox, QSpinBox, QComboBox, QMessageBox,
    QTabWidget, QTextEdit, QGroupBox
)
from PyQt6.QtCore import Qt, pyqtSignal, QThread
from PyQt6.QtGui import QIcon, QColor, QFont


# 配置Tesseract路径（Windows）
def setup_tesseract():
    """设置Tesseract和语言包路径"""
    if sys.platform == 'win32':
        tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        tessdata_path = r'C:\Program Files\Tesseract-OCR\tessdata'
        
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.pytesseract_cmd = tesseract_path
        
        if os.path.exists(tessdata_path):
            os.environ['TESSDATA_PREFIX'] = tessdata_path


def download_language_pack(lang_code: str, tessdata_path: str) -> bool:
    """下载Tesseract语言包"""
    try:
        os.makedirs(tessdata_path, exist_ok=True)
        
        lang_file = os.path.join(tessdata_path, f"{lang_code}.traineddata")
        if os.path.exists(lang_file):
            return True
        
        print(f"正在下载语言包: {lang_code}...")
        
        # 尝试多个源
        urls = [
            f"https://github.com/UB-Mannheim/tesseract/raw/master/tessdata/{lang_code}.traineddata",
            f"https://raw.githubusercontent.com/UB-Mannheim/tesseract/master/tessdata/{lang_code}.traineddata",
            f"https://github.com/tesseract-ocr/tessdata/raw/main/{lang_code}.traineddata",
        ]
        
        last_error = None
        for url in urls:
            try:
                print(f"  尝试下载源: {url}")
                urllib.request.urlretrieve(url, lang_file)
                print(f"✓ 语言包下载成功: {lang_file}")
                return True
            except Exception as e:
                last_error = str(e)
                print(f"  此源下载失败，尝试下一个...")
                continue
        
        print(f"✗ 所有源均下载失败: {last_error}")
        return False
            
    except Exception as e:
        print(f"✗ 错误: {str(e)}")
        return False


class ConversionWorker(QThread):
    """后台转换线程"""
    progress = pyqtSignal(int)
    message = pyqtSignal(str)
    finished = pyqtSignal(bool)
    
    def __init__(self, files: List[str], output_dir: str, use_ocr: bool, 
                 ocr_lang: str = 'chi_sim', dpi: int = 200):
        super().__init__()
        self.files = files
        self.output_dir = output_dir
        self.use_ocr = use_ocr
        self.ocr_lang = ocr_lang
        self.dpi = dpi
        self.tessdata_path = r'C:\Program Files\Tesseract-OCR\tessdata'
        
    def run(self):
        """执行转换"""
        try:
            total_files = len(self.files)
            
            for index, pdf_file in enumerate(self.files):
                try:
                    self.message.emit(f"正在处理: {os.path.basename(pdf_file)}")
                    
                    # 检查文件是否存在
                    if not os.path.exists(pdf_file):
                        self.message.emit(f"✗ 错误: 文件不存在 - {pdf_file}")
                        progress_percent = int((index + 1) / total_files * 100)
                        self.progress.emit(progress_percent)
                        continue
                    
                    # 生成输出文件名
                    output_path = os.path.join(
                        self.output_dir,
                        Path(pdf_file).stem + '.docx'
                    )
                    
                    # 如果启用了OCR，优先使用OCR转换
                    if self.use_ocr:
                        self._convert_with_ocr(pdf_file, output_path)
                    else:
                        # 先尝试普通转换，失败则自动尝试OCR转换
                        try:
                            self._convert_normal(pdf_file, output_path)
                        except Exception as e:
                            self.message.emit(f"  ⚠ 普通转换失败，尝试OCR转换...")
                            self._convert_with_ocr(pdf_file, output_path)
                    
                    # 检查输出文件是否成功创建
                    if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                        self.message.emit(f"✓ 完成: {os.path.basename(pdf_file)}")
                    else:
                        self.message.emit(f"✗ 错误: 输出文件生成失败")
                    
                    progress_percent = int((index + 1) / total_files * 100)
                    self.progress.emit(progress_percent)
                    
                except Exception as e:
                    self.message.emit(f"✗ 错误 ({os.path.basename(pdf_file)}): {str(e)}")
                    progress_percent = int((index + 1) / total_files * 100)
                    self.progress.emit(progress_percent)
            
            self.finished.emit(True)
            
        except Exception as e:
            self.message.emit(f"转换失败: {str(e)}")
            self.finished.emit(False)
    
    def _convert_normal(self, pdf_file: str, output_path: str):
        """普通转换（适用于数字PDF）"""
        try:
            cv = Converter(pdf_file)
            cv.convert(output_path, start=0, end=None)
            cv.close()
        except Exception as e:
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
            raise Exception(f"PDF转换失败: {str(e)}")
    
    def _convert_with_ocr(self, pdf_file: str, output_path: str):
        """使用OCR转换扫描版PDF"""
        temp_dir = None
        try:
            # 检查和下载语言包
            lang_file = os.path.join(self.tessdata_path, f"{self.ocr_lang}.traineddata")
            if not os.path.exists(lang_file):
                self.message.emit(f"  → 准备OCR语言包: {self.ocr_lang}...")
                if not download_language_pack(self.ocr_lang, self.tessdata_path):
                    self.message.emit(f"  ⚠ 语言包下载失败，尝试使用英文...")
                    # 降级到英文尝试
                    if self.ocr_lang != 'eng':
                        eng_file = os.path.join(self.tessdata_path, "eng.traineddata")
                        if os.path.exists(eng_file):
                            self.message.emit(f"  ✓ 使用英文OCR")
                            self.ocr_lang = 'eng'
                        else:
                            raise Exception(f"无法获取OCR语言包，请检查Tesseract安装")
                    else:
                        raise Exception(f"英文语言包也不可用")
                else:
                    self.message.emit(f"  ✓ 语言包已准备")
            
            # 创建临时目录存放图片
            temp_dir = tempfile.mkdtemp()
            
            # 将PDF转换为图片
            self.message.emit(f"  → 提取PDF页面中...")
            try:
                images = convert_from_path(pdf_file, dpi=self.dpi)
            except Exception as e:
                self.message.emit(f"  ✗ PDF提取失败: {str(e)}")
                raise Exception(f"无法提取PDF页面: {str(e)}")
            
            if not images:
                raise Exception("PDF页面提取失败，可能是文件损坏")
            
            self.message.emit(f"  → 共提取 {len(images)} 页")
            
            # 创建新的Word文档
            doc = Document()
            
            total_pages = len(images)
            successful_pages = 0
            
            for page_num, image in enumerate(images, 1):
                try:
                    self.message.emit(f"  → OCR识别第 {page_num}/{total_pages} 页...")
                    
                    # 使用Tesseract进行OCR识别
                    text = pytesseract.image_to_string(image, lang=self.ocr_lang)
                    
                    # 保存图片到临时目录
                    temp_image_path = os.path.join(temp_dir, f"page_{page_num}.png")
                    image.save(temp_image_path)
                    
                    # 添加图片
                    try:
                        doc.add_picture(temp_image_path, width=Inches(6))
                    except Exception as e:
                        self.message.emit(f"  ⚠ 无法添加第{page_num}页图片")
                    
                    # 添加识别的文本
                    if text and text.strip():
                        p = doc.add_paragraph(text)
                        p.style = 'Normal'
                        successful_pages += 1
                    
                    # 添加页分符（除了最后一页）
                    if page_num < total_pages:
                        doc.add_page_break()
                        
                except Exception as e:
                    self.message.emit(f"  ⚠ 第{page_num}页OCR失败: {str(e)}")
                    continue
            
            # 检查是否至少成功处理了一些页面
            if successful_pages == 0:
                self.message.emit(f"  ⚠ 未能识别任何文本，仅保存图片")
            else:
                self.message.emit(f"  ✓ 成功识别 {successful_pages}/{total_pages} 页")
            
            # 保存文档
            doc.save(output_path)
            self.message.emit(f"  → 文档已保存")
            
        except Exception as e:
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
            self.message.emit(f"  ✗ OCR转换失败: {str(e)}")
            raise Exception(f"OCR转换失败: {str(e)}")
        
        finally:
            # 清理临时文件
            if temp_dir and os.path.exists(temp_dir):
                try:
                    import shutil
                    shutil.rmtree(temp_dir)
                except:
                    pass


class PDFConverterUI(QMainWindow):
    """PDF转Word转换器UI"""
    
    def __init__(self):
        super().__init__()
        self.pdf_files: List[str] = []
        self.worker = None
        self.convert_btn = None
        setup_tesseract()  # 初始化Tesseract设置
        self.init_ui()
        
    def init_ui(self):
        """初始化UI"""
        self.setWindowTitle("PDF转Word转换器 v1.0")
        self.setGeometry(100, 100, 900, 700)
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QVBoxLayout()
        
        # 创建标题
        title = QLabel("PDF转Word批量转换工具")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title.setFont(title_font)
        main_layout.addWidget(title)
        
        # 创建标签页
        tabs = QTabWidget()
        main_layout.addWidget(tabs)
        
        # 转换标签页
        convert_widget = self._create_convert_tab()
        tabs.addTab(convert_widget, "转换")
        
        # 设置标签页
        settings_widget = self._create_settings_tab()
        tabs.addTab(settings_widget, "设置")
        
        # 日志标签页
        log_widget = self._create_log_tab()
        tabs.addTab(log_widget, "日志")
        
        central_widget.setLayout(main_layout)
    
    def _create_convert_tab(self) -> QWidget:
        """创建转换标签页"""
        widget = QWidget()
        layout = QVBoxLayout()
        
        # 文件列表部分
        file_group = QGroupBox("选择PDF文件")
        file_layout = QVBoxLayout()
        
        # 文件列表
        self.file_list = QListWidget()
        self.file_list.setMaximumHeight(250)
        file_layout.addWidget(QLabel("已选择的文件:"))
        file_layout.addWidget(self.file_list)
        
        # 文件操作按钮
        button_layout = QHBoxLayout()
        
        add_btn = QPushButton("添加文件")
        add_btn.clicked.connect(self.add_files)
        button_layout.addWidget(add_btn)
        
        add_folder_btn = QPushButton("添加文件夹")
        add_folder_btn.clicked.connect(self.add_folder)
        button_layout.addWidget(add_folder_btn)
        
        remove_btn = QPushButton("移除选中")
        remove_btn.clicked.connect(self.remove_file)
        button_layout.addWidget(remove_btn)
        
        clear_btn = QPushButton("清空列表")
        clear_btn.clicked.connect(self.clear_files)
        button_layout.addWidget(clear_btn)
        
        file_layout.addLayout(button_layout)
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        # 选项部分
        options_group = QGroupBox("转换选项")
        options_layout = QVBoxLayout()
        
        # OCR选项
        self.ocr_checkbox = QCheckBox("使用OCR识别扫描版PDF (推荐扫描版PDF使用)")
        self.ocr_checkbox.setChecked(False)
        self.ocr_checkbox.stateChanged.connect(self.on_ocr_changed)
        options_layout.addWidget(self.ocr_checkbox)
        
        # OCR语言选择
        lang_layout = QHBoxLayout()
        lang_layout.addWidget(QLabel("OCR语言:"))
        self.ocr_lang_combo = QComboBox()
        self.ocr_lang_combo.addItems(['中文简体 (chi_sim)', '中文繁体 (chi_tra)', 
                                      '英文 (eng)', '日文 (jpn)'])
        self.ocr_lang_combo.setEnabled(False)
        lang_layout.addWidget(self.ocr_lang_combo)
        lang_layout.addStretch()
        options_layout.addLayout(lang_layout)
        
        # DPI设置
        dpi_layout = QHBoxLayout()
        dpi_layout.addWidget(QLabel("DPI (仅OCR):"))
        self.dpi_spinbox = QSpinBox()
        self.dpi_spinbox.setValue(200)
        self.dpi_spinbox.setRange(100, 600)
        self.dpi_spinbox.setEnabled(False)
        dpi_layout.addWidget(self.dpi_spinbox)
        dpi_layout.addStretch()
        options_layout.addLayout(dpi_layout)
        
        # 输出目录
        output_layout = QHBoxLayout()
        output_layout.addWidget(QLabel("输出目录:"))
        self.output_label = QLabel(str(Path.home() / "Desktop"))
        output_layout.addWidget(self.output_label)
        
        output_btn = QPushButton("浏览...")
        output_btn.clicked.connect(self.choose_output_dir)
        output_layout.addWidget(output_btn)
        options_layout.addLayout(output_layout)
        
        options_group.setLayout(options_layout)
        layout.addWidget(options_group)
        
        # 进度部分
        progress_group = QGroupBox("进度")
        progress_layout = QVBoxLayout()
        
        self.progress_label = QLabel("准备就绪")
        progress_layout.addWidget(self.progress_label)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        progress_layout.addWidget(self.progress_bar)
        
        progress_group.setLayout(progress_layout)
        layout.addWidget(progress_group)
        
        # 控制按钮
        control_layout = QHBoxLayout()
        control_layout.addStretch()
        
        self.convert_btn = QPushButton("开始转换")
        self.convert_btn.setMinimumHeight(40)
        self.convert_btn.setMinimumWidth(120)
        self.convert_btn.clicked.connect(self.start_conversion)
        control_layout.addWidget(self.convert_btn)
        
        layout.addLayout(control_layout)
        
        widget.setLayout(layout)
        return widget
    
    def _create_settings_tab(self) -> QWidget:
        """创建设置标签页"""
        widget = QWidget()
        layout = QVBoxLayout()
        
        info_text = QTextEdit()
        info_text.setReadOnly(True)
        info_text.setText("""
<h3>关于本工具</h3>
<p><b>PDF转Word批量转换工具 v1.0</b></p>

<h4>功能特性:</h4>
<ul>
  <li>✓ 支持批量转换PDF文件</li>
  <li>✓ 支持普通数字PDF和扫描版PDF</li>
  <li>✓ 支持OCR文字识别（扫描版）</li>
  <li>✓ 支持多种语言识别（中文、英文、日文等）</li>
  <li>✓ 自动下载OCR语言包</li>
  <li>✓ 友好的图形化界面</li>
  <li>✓ 实时转换日志显示</li>
  <li>✓ 自动错误恢复机制</li>
</ul>

<h4>使用步骤:</h4>
<ol>
  <li>在"转换"标签页添加要转换的PDF文件</li>
  <li>对于扫描版PDF，建议勾选"使用OCR识别"选项</li>
  <li>对于普通数字PDF，可以不勾选OCR（速度更快）</li>
  <li>设置输出目录</li>
  <li>点击"开始转换"按钮</li>
</ol>

<h4>转换模式:</h4>
<p><b>数字PDF（默认）:</b></p>
<ul>
  <li>• 使用pdf2docx库直接转换</li>
  <li>• 速度快，保留排版</li>
  <li>• 适合纯文本PDF</li>
</ul>

<p><b>扫描版PDF（OCR）:</b></p>
<ul>
  <li>• 先转图片，再用Tesseract OCR识别</li>
  <li>• 自动下载所需语言包</li>
  <li>• 处理时间较长</li>
  <li>• 适合扫描件、图片PDF</li>
  <li>• 无法识别时仅保存图片</li>
  <li>• 如果语言包下载失败，会自动降级使用英文</li>
</ul>

<h4>系统要求:</h4>
<ul>
  <li>• Python 3.7+</li>
  <li>• Tesseract-OCR (自动配置和下载)</li>
  <li>• 网络连接（用于下载语言包）</li>
</ul>

<h4>常见问题:</h4>
<p><b>Q: 语言包下载失败怎么办？</b></p>
<ul>
  <li>• 检查网络连接</li>
  <li>• 程序会自动降级使用英文OCR</li>
  <li>• 如果需要中文，请手动下载语言包到 C:\\Program Files\\Tesseract-OCR\\tessdata</li>
</ul>

<p><b>Q: 如何手动下载语言包？</b></p>
<ul>
  <li>• 从以下链接下载: https://github.com/tesseract-ocr/tessdata/tree/main</li>
  <li>• 放在: C:\\Program Files\\Tesseract-OCR\\tessdata</li>
  <li>• 例如: chi_sim.traineddata（中文简体）</li>
</ul>
        """)
        layout.addWidget(info_text)
        
        widget.setLayout(layout)
        return widget
    
    def _create_log_tab(self) -> QWidget:
        """创建日志标签页"""
        widget = QWidget()
        layout = QVBoxLayout()
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        layout.addWidget(self.log_text)
        
        clear_log_btn = QPushButton("清空日志")
        clear_log_btn.clicked.connect(lambda: self.log_text.clear())
        layout.addWidget(clear_log_btn)
        
        widget.setLayout(layout)
        return widget
    
    def add_files(self):
        """添加PDF文件"""
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择PDF文件", "",
            "PDF文件 (*.pdf);;所有文件 (*.*)"
        )
        for f in files:
            if f not in self.pdf_files:
                self.pdf_files.append(f)
        self.update_file_list()
    
    def add_folder(self):
        """添加文件夹中的所有PDF"""
        folder = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if folder:
            for f in Path(folder).glob("*.pdf"):
                if str(f) not in self.pdf_files:
                    self.pdf_files.append(str(f))
            self.update_file_list()
    
    def remove_file(self):
        """移除选中的文件"""
        for item in self.file_list.selectedItems():
            self.pdf_files.remove(item.text())
        self.update_file_list()
    
    def clear_files(self):
        """清空文件列表"""
        reply = QMessageBox.question(
            self, "确认", "确定要清空所有文件吗?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self.pdf_files.clear()
            self.update_file_list()
    
    def update_file_list(self):
        """更新文件列表显示"""
        self.file_list.clear()
        for f in self.pdf_files:
            item = QListWidgetItem(f)
            self.file_list.addItem(item)
    
    def on_ocr_changed(self):
        """OCR复选框状态改变"""
        enabled = self.ocr_checkbox.isChecked()
        self.ocr_lang_combo.setEnabled(enabled)
        self.dpi_spinbox.setEnabled(enabled)
    
    def choose_output_dir(self):
        """选择输出目录"""
        folder = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if folder:
            self.output_label.setText(folder)
    
    def start_conversion(self):
        """开始转换"""
        if not self.pdf_files:
            QMessageBox.warning(self, "警告", "请先添加PDF文件!")
            return
        
        output_dir = self.output_label.text()
        if not os.path.exists(output_dir):
            QMessageBox.warning(self, "错误", "输出目录不存在!")
            return
        
        # 禁用按钮
        self.convert_btn.setEnabled(False)
        
        # 获取OCR选项
        use_ocr = self.ocr_checkbox.isChecked()
        ocr_lang_map = {
            '中文简体 (chi_sim)': 'chi_sim',
            '中文繁体 (chi_tra)': 'chi_tra',
            '英文 (eng)': 'eng',
            '日文 (jpn)': 'jpn'
        }
        ocr_lang = ocr_lang_map.get(self.ocr_lang_combo.currentText(), 'chi_sim')
        dpi = self.dpi_spinbox.value()
        
        # 创建工作线程
        self.worker = ConversionWorker(
            self.pdf_files,
            output_dir,
            use_ocr,
            ocr_lang,
            dpi
        )
        self.worker.progress.connect(self.on_progress)
        self.worker.message.connect(self.on_message)
        self.worker.finished.connect(self.on_finished)
        self.worker.start()
        
        self.progress_label.setText("正在转换...")
        self.log_text.append(f"[{datetime.now().strftime('%H:%M:%S')}] 开始转换 {len(self.pdf_files)} 个文件")
        self.log_text.append(f"[{datetime.now().strftime('%H:%M:%S')}] 输出目录: {output_dir}")
        self.log_text.append(f"[{datetime.now().strftime('%H:%M:%S')}] OCR: {'已启用' if use_ocr else '已禁用'}\n")
    
    def on_progress(self, value):
        """进度更新"""
        self.progress_bar.setValue(value)
    
    def on_message(self, message):
        """消息更新"""
        self.progress_label.setText(message)
        self.log_text.append(f"[{datetime.now().strftime('%H:%M:%S')}] {message}")
    
    def on_finished(self, success):
        """转换完成"""
        self.convert_btn.setEnabled(True)
        if success:
            self.progress_label.setText("转换完成!")
            output_dir = self.output_label.text()
            QMessageBox.information(
                self, "成功",
                f"转换完成！\n\n输出目录:\n{output_dir}\n\n请在文件管理器中查看结果。"
            )
        else:
            self.progress_label.setText("转换失败")
            QMessageBox.critical(self, "错误", "转换过程中出现错误，请查看日志了解详情。")


def main():
    """主函数"""
    app = QApplication(sys.argv)
    window = PDFConverterUI()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
